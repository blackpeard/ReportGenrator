"""
report_engine.py
================
Centralized report generation engine.
Shared by GUI, CLI, and any future interface.
"""

from __future__ import annotations

import os
from pydoc import doc
import re
import sys
import shutil
import zipfile
from lxml import etree
from docx.oxml.ns import qn
import openpyxl

# ── MUST be first — anchor all paths to this file's location ────────────────
if getattr(sys, 'frozen', False):
    _ENGINE_DIR = sys._MEIPASS          # running as .exe
else:
    _ENGINE_DIR = os.path.dirname(os.path.abspath(__file__))

# ── src/ imports anchored — works from any CWD or subfolder ─────────────────
_SRC = os.path.join(_ENGINE_DIR, "src")
if _SRC not in sys.path:
    sys.path.insert(0, _SRC)

from dataclasses import dataclass, field
from datetime import datetime
from typing import Callable, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate


from excel_reader import ExcelReader
from poc_finder import POCFinder


# ── Data classes ──────────────────────────────────────────────────────────────

@dataclass
class ReportConfig:
    """
    All inputs needed to generate one report.

    Data source modes
    -----------------
    Excel mode  : set excel_file path, leave manual_observations empty
    Manual mode : set manual_observations list, leave excel_file empty
    """

    # ── Report / client info ─────────────────────────────────────────────────
    client_name:    str = ""
    app_name:       str = ""
    app_type:       str = ""
    audit_period:   str = ""
    url:            str = ""
    method:         str = ""
    report_type:    str = "Web"         # Web | Api | Mobile
    environment:    str = "Production"  # Production | Uat

    # ── Cover / front-page fields ────────────────────────────────────────────
    prepared_by:                str = ""
    prepared_by_designation:    str = ""
    reviewed_by:                str = ""
    reviewed_by_designation:    str = ""
    doc_version:                str = "1.0"
    doc_status:                 str = "Draft"       # Draft | Final
    client_history:             str = ""
    limitation:                 str = ""
    # Document approval
    approved_by:                str = ""
    approved_by_designation:    str = ""
    released_by:                str = ""
    released_by_designation:    str = ""
    release_date:               str = ""

    # Client contact
    client_contact_person:  str = ""
    client_designation:     str = ""
    client_email:           str = ""

    # Tools used — list of dicts with keys:
    # tool_name, tool_version, tool_type, category
    selected_tools: list = field(default_factory=list)
    team_members: list = field(default_factory=list)
    # keys per dict: name, designation, email, qualifications, cert_in_listed

    # ── Data source — only ONE should be set at a time ───────────────────────
    excel_file:             str  = ""
    manual_observations:    list = field(default_factory=list)
    # manual_observations keys per dict:
    #   sr_no, title, severity, description, impact,
    #   recommendation, affected_url, cve

    # ── Files ────────────────────────────────────────────────────────────────
    poc_folder:     str = ""
    output_file:    str = ""
    template_base:  str = ""   # blank = auto → _ENGINE_DIR/templates/

    # ── Optional count overrides (blank = auto-calculate) ────────────────────
    critical_count: str = ""
    high_count:     str = ""
    medium_count:   str = ""
    low_count:      str = ""
    total_count:    str = ""

    # ── Scope (for Word report) ──────────────────────────────────────────────
    scope: list = field(default_factory=list)  



    def template_path(self) -> str:
        mapping = {
            "web":    "web_template.docx",
            "api":    "api_template.docx",
            "android": "android_template.docx",
            "ios":     "ios_template.docx",
            "va":      "va_template.docx",
            "ca":      "ca_template.docx",
            "ca_nessus": "ca_nessus_template.docx",
            "sourcecode": "sourcecode_template.docx",
        }
        filename = mapping.get(self.report_type.lower(), "web_template.docx")
        base = self.template_base or os.path.join(_ENGINE_DIR, "templates")
        return os.path.join(base, filename)

    def resolved_output(self, safe_name: str) -> str:
            if self.output_file:
                p = self.output_file
                return p if p.endswith(".docx") else p + ".docx"
            stamp  = datetime.now().strftime("%Y%m%d")
            prefix = "Final_Report" if self.doc_status.lower() == "final" else "Draft_Report"
            return f"{prefix}_{safe_name}_{stamp}.docx"

    @property
    def is_manual_mode(self) -> bool:
        """True when observations supplied directly — Excel not needed."""
        return bool(self.manual_observations) and not self.excel_file


# ── Batch Runner ──────────────────────────────────────────────────────────────

@dataclass
class BatchResult:
    total:    int = 0
    success:  int = 0
    failed:   int = 0
    errors:   list = field(default_factory=list)  # list of (filename, error)


class BatchRunner:
    """
    Run ReportEngine on every .xlsx file in a folder using the same config.

    Usage:
        runner = BatchRunner(
            base_config   = config,        # ReportConfig — excel_file left blank
            excel_folder  = "C:/excels/",
            output_folder = "C:/reports/",
            progress_callback = fn,        # optional — called per file (0-100)
            log_callback      = fn,        # optional
        )
        result = runner.run()
    """

    def __init__(
        self,
        base_config:       ReportConfig,
        excel_folder:      str,
        output_folder:     str,
        progress_callback: Optional[Callable[[int], None]] = None,
        log_callback:      Optional[Callable[[str], None]] = None,
       
    ):
        self.base_config    = base_config
        self.excel_folder   = excel_folder
        self.output_folder  = output_folder
        self._progress      = progress_callback or (lambda _: None)
        self._log           = log_callback or print
      

    def run(self) -> BatchResult:
        import copy

        result = BatchResult()
        os.makedirs(self.output_folder, exist_ok=True)

        # Collect all xlsx files
        excel_files = [
            f for f in os.listdir(self.excel_folder)
            if f.lower().endswith((".xlsx", ".xls"))
        ]

        if not excel_files:
            self._log("⚠️  No Excel files found in folder.")
            return result

        result.total = len(excel_files)
        self._log(f"📂 Found {result.total} Excel file(s) — starting batch...")

        for i, filename in enumerate(excel_files, 1):
            self._log(f"\n{'─'*40}")
            self._log(f"📄 [{i}/{result.total}] Processing: {filename}")

            # Deep copy config so each run is independent
            cfg = copy.deepcopy(self.base_config)
            cfg.excel_file  = os.path.join(self.excel_folder, filename)
            status = getattr(cfg, 'doc_status', 'Draft')
            prefix = "Final_Report" if status.lower() == "final" else "Draft_Report"
            cfg.output_file = os.path.join(
                self.output_folder,
                f"{prefix}_{os.path.splitext(filename)[0]}.docx"
            )

            def _per_file_progress(pct, idx=i, total=result.total):
                # Scale per-file 0-100 into overall batch progress
                overall = int(((idx - 1) / total * 100) + (pct / total))
                self._progress(overall)

            engine = ReportEngine(
                cfg,
                progress_callback=_per_file_progress,
                log_callback=self._log,
            )
            res = engine.run()

            if res.success:
                result.success += 1
                self._log(f"✅ Done: {res.output_path}")
            else:
                result.failed += 1
                result.errors.append((filename, res.error))
                self._log(f"❌ Failed: {filename} — {res.error}")

        self._progress(100)
        self._log(f"\n{'═'*40}")
        self._log(f"📊 Batch complete — ✅ {result.success} succeeded  ❌ {result.failed} failed")
        if result.errors:
            for fname, err in result.errors:
                self._log(f"   ✗ {fname}: {err}")

        return result

@dataclass
class ReportResult:
    success:            bool
    output_path:        str = ""
    error:              str = ""
    observations_count: int = 0


# ── Text utilities ────────────────────────────────────────────────────────────

def _clean_text(text) -> str:
    if not isinstance(text, str):
        text = str(text) if text is not None else ""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text


def _clean_observations(observations: list[dict]) -> list[dict]:
    return [{k: _clean_text(v) for k, v in obs.items()} for obs in observations]


def _auto_number(observations: list[dict]) -> list[dict]:
    for i, obs in enumerate(observations, 1):
        if not obs.get("sr_no"):
            obs["sr_no"] = str(i)
    return observations


# ── Engine ────────────────────────────────────────────────────────────────────

class ReportEngine:
    """
    Single reusable engine. Import and call engine.run() from anywhere.

    Parameters
    ----------
    config            : ReportConfig
    progress_callback : callable(int 0-100)
    log_callback      : callable(str)
    """

    SEVERITY_COLORS = {
        "critical": "C00000",
        "high":     "EE0000",
        "medium":   "FFC000",
        "low":      "00B050",
        "info":     "0070C0",
    }

    def __init__(
        self,
        config:            ReportConfig,
        progress_callback: Optional[Callable[[int], None]] = None,
        log_callback:      Optional[Callable[[str], None]] = None,

        
    ):
        self.config    = config
        self._progress = progress_callback or (lambda _: None)
        self._log      = log_callback or print
       

    def run(self) -> ReportResult:
        try:
            return self._execute()
        except Exception as exc:
            import traceback
            self._log(f"❌ Fatal error: {exc}")
            self._log(traceback.format_exc())
            return ReportResult(success=False, error=str(exc))

    def _execute(self) -> ReportResult:
        cfg = self.config
        self._progress(0)

        # ── 1. Validate ───────────────────────────────────────────────────────
        self._log("🔍 Validating inputs...")
        if cfg.is_manual_mode:
            if not cfg.manual_observations:
                raise ValueError("Manual mode: no observations provided.")
        else:
            if not cfg.excel_file or not os.path.exists(cfg.excel_file):
                raise FileNotFoundError(f"Excel file not found: {cfg.excel_file}")

        if not os.path.exists(cfg.template_path()):
            raise FileNotFoundError(f"Template not found: {cfg.template_path()}")
        self._progress(10)

        # ── 2. Load observations ──────────────────────────────────────────────
        scope_text      = ""
        limitation_text = cfg.limitation or "No limitations specified"
        report_name     = cfg.app_name or "Report"
        stats           = {"critical": "0", "high": "0", "medium": "0", "low": "0"}

        if cfg.scope:
            scope_text = "\n".join(cfg.scope) if isinstance(cfg.scope, list) else str(cfg.scope)

        if cfg.is_manual_mode:
            self._log("📋 Using manual observations...")
            observations = list(cfg.manual_observations)
            safe_name    = re.sub(r'[\\/*?:"<>|]', "", cfg.app_name or "Report")

        else:
            self._log("📊 Loading Excel file...")
            reader = ExcelReader()
            reader.load(cfg.excel_file)
            
            data = reader.read_all()
            
            observations = data['observations']
    
        
            excel_scope = data.get('scope', [])
            limitation = data['limitation']
            
            
            index_df = reader.read_index_sheet()
            report_name = reader.extract_report_name(index_df)
            safe_name = re.sub(r'[\\/*?:"<>|]', "", report_name)
            stats = reader.extract_summary_stats(index_df)

            if cfg.scope:
                scope_text = "\n".join(cfg.scope) if isinstance(cfg.scope, list) else str(cfg.scope)
            else:
                scope_text = "\n".join(excel_scope) if excel_scope else "Full application scope"
                        

            limitation_text = "\n".join(limitation) if limitation else "No limitations specified"

        self._progress(25)

        # ── 3. Clean + number ─────────────────────────────────────────────────
        self._log("🧹 Cleaning data...")
        observations = _clean_observations(observations)
        observations = _auto_number(observations)
        self._progress(35)

        # ── 4. Severity counts ────────────────────────────────────────────────
        critical = cfg.critical_count or self._count_sev(observations, ["critical"])     or stats.get("critical", "0")
        high   = cfg.high_count   or self._count_sev(observations, ["high"])             or stats.get("high",   "0")
        medium = cfg.medium_count or self._count_sev(observations, ["medium"])           or stats.get("medium", "0")
        low    = cfg.low_count    or self._count_sev(observations, ["low"])              or stats.get("low",    "0")
        total  = cfg.total_count  or str(len(observations))
        self._log(f"📊 Critical: {critical} High: {high}  Medium: {medium}  Low: {low}  Total: {total}")

        # ── 5. POC scanner ────────────────────────────────────────────────────
      
        poc_finder = POCFinder()

        excel_pocs_loaded = False
        if not cfg.is_manual_mode:
            try:
                excel_pocs_loaded = poc_finder.load_pocs_from_excel(cfg.excel_file)
            except Exception as e:
                self._log(f"⚠️  Excel POC load failed: {e}")

        if excel_pocs_loaded:
            self._log("✅ Using POCs from Excel POC sheet")
            for obs in observations:
                vuln_title = obs.get("title", "").strip()
                ordered_items = poc_finder.get_excel_poc_items_by_vulnerability(vuln_title)
                obs['_poc_items'] = ordered_items
                if ordered_items:
                    self._log(f"    ✓ {len(ordered_items)} POC items matched → {vuln_title[:40]}")
                else:
                    self._log(f"    ⚠️  No POC items matched → {vuln_title[:40]}")
        else:
            # Fallback — folder-based POC finder
            self._log("📁 Excel POC sheet not found — trying POC folder...")
            if cfg.poc_folder and os.path.exists(cfg.poc_folder):
                poc_finder.scan_folder(cfg.poc_folder)
                self._log(f"🖼️  POC folders found: {len(poc_finder.get_all_vulnerability_folders())}")
            elif cfg.poc_folder:
                self._log(f"⚠️  POC folder not found: {cfg.poc_folder}")
            else:
                self._log("ℹ️  No POC source provided — skipping POC insertion")

        # ── 6. Build context ──────────────────────────────────────────────────
        context = {
            "App_Name":     _clean_text(cfg.app_name),
            "type":         _clean_text(cfg.app_type),
            "Audit_Period": _clean_text(cfg.audit_period),
            "Client_Name":  _clean_text(cfg.client_name),
            "environment":  cfg.environment.capitalize(),
            "URL":          _clean_text(cfg.url),
            "method":       _clean_text(cfg.method),
            "report_name":  report_name,
            "date_today":   datetime.today().strftime("%d %B %Y"),
            "observations": observations,
            "scope":        scope_text,
            "limitation":   _clean_text(limitation_text),
            "critical":     critical,
            "high":         high,
            "medium":       medium,
            "low":          low,
            "total":        total,
            # Cover fields — use {{ prepared_by }}, {{ doc_version }} etc. in .docx template
            "prepared_by":              _clean_text(cfg.prepared_by),
            "prepared_by_designation":  _clean_text(cfg.prepared_by_designation),
            "reviewed_by":              _clean_text(cfg.reviewed_by),
            "reviewed_by_designation":  _clean_text(cfg.reviewed_by_designation),
            "doc_version":              _clean_text(cfg.doc_version),
            "doc_status":               _clean_text(cfg.doc_status),
            "client_history":           _clean_text(cfg.client_history),

            "approved_by":              _clean_text(cfg.approved_by),
            "approved_by_designation":  _clean_text(cfg.approved_by_designation),
            "released_by":              _clean_text(cfg.released_by),
            "released_by_designation":  _clean_text(cfg.released_by_designation),
            "release_date":             _clean_text(cfg.release_date),
            "client_contact_person":    _clean_text(cfg.client_contact_person),
            "client_designation":       _clean_text(cfg.client_designation),
            "client_email":             _clean_text(cfg.client_email),
           
        }
        self._progress(55)

        # ── 7. Render template ────────────────────────────────────────────────
        # self._log("📝 Rendering template...")
        # tpl       = DocxTemplate(cfg.template_path())
        # tpl.render(context)
        # for para in tpl.docx.paragraphs:
        #     text = para.text.strip()
        #     if text.startswith("{%") and text.endswith("%}"):
        #         para.paragraph_format.space_before = Pt(0)
        #         para.paragraph_format.space_after  = Pt(0)
        # temp_path = os.path.join(_ENGINE_DIR, "_render_temp.docx")
        # tpl.save(temp_path)
        # self._progress(65)
        self._log("📝 Rendering template...")
        from jinja2 import Environment
        from jinja2 import Environment

        tpl = DocxTemplate(cfg.template_path())
        jinja_env = Environment(trim_blocks=True, lstrip_blocks=True)
        tpl.render(context, jinja_env)
        temp_path = os.path.join(_ENGINE_DIR, "_render_temp.docx")
        tpl.save(temp_path)

        # ── 8. Post-process ───────────────────────────────────────────────────
        self._log("🎨 Building executive summary table...")
        doc = Document(temp_path)
        self._build_exec_table(doc, observations, cfg.report_type.lower())
        self._insert_page_breaks(doc, observations)
        # self._remove_jinja_artifact_paragraphs(doc)
        self._remove_empty_rows(doc)

        # Build auditing team table (6-column)
        self._log(f"DEBUG — team_members count: {len(cfg.team_members)}")
        self._log(f"DEBUG — selected_tools count: {len(cfg.selected_tools)}")

        if cfg.team_members:
            self._log("👥 Building auditing team table...")
            self._build_team_table(doc, cfg.team_members)
        else:
            self._log("⚠️  No team members received — skipping team table.")

        if cfg.selected_tools:
            self._log("🔧 Building tools table...")
            self._build_tools_table(doc, cfg.selected_tools)
        else:
            self._log("⚠️  No tools received — skipping tools table.")
        
        if cfg.scope:
            self._log("📋 Building scope table...")
            self._build_scope_table(doc, cfg.scope, cfg.app_type, cfg.environment)

        # ── 9. POC images ─────────────────────────────────────────────────────
        has_excel_pocs  = excel_pocs_loaded
        has_folder_pocs = bool(cfg.poc_folder and os.path.exists(cfg.poc_folder))

        if (has_excel_pocs or has_folder_pocs) and observations:
            self._log("🖼️  Inserting POC screenshots...")
            self._log(f"DEBUG has_excel_pocs={has_excel_pocs} has_folder_pocs={has_folder_pocs}")
            self._log(f"DEBUG observations count={len(observations)}")
            for idx, obs in enumerate(observations):
                self._log(f"DEBUG obs[{idx}] title='{obs.get('title','')}' _poc_images={obs.get('_poc_images', [])}")
            # Check if marker exists in document
            marker = "<!-- POC will be inserted here during post-processing -->"
            marker_count = sum(1 for p in doc.paragraphs if marker in p.text)
            self._log(f"DEBUG marker found in {marker_count} paragraphs")
            self._insert_pocs(doc, observations, poc_finder)

        # ── 10. Save ──────────────────────────────────────────────────────────
        output_path = cfg.resolved_output(safe_name)
        doc.save(output_path)
        self._log("📈 Updating chart data...")
        self._update_chart_data(output_path, critical, high, medium, low)
        try:
            os.remove(temp_path)
        except OSError:
            pass
        self._progress(100)

        self._log(f"✅ Report saved: {output_path}")
        self._log("⚠️  Please review the report carefully before sharing.")

        return ReportResult(
            success=True,
            output_path=output_path,
            observations_count=len(observations),
        )

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _count_sev(observations: list[dict], labels: list[str]) -> str:
        n = sum(1 for o in observations if o.get("severity", "").lower() in labels)
        return str(n) if n else ""

    def _build_exec_table(self, doc: Document, observations: list[dict], template_type: str = "web"):
        """Build executive summary table using template type from GUI."""
        
        # Define field mappings for each template type
        FIELD_MAPPING = {
            "ca_nessus": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3, "checkid": 4, "recommendation": 5, "new_obs": 6},
            "web": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3, "cve": 4, "recommendation": 5, "new_obs": 6},
            "va": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3, "recommendation": 4, "new_obs": 5},
            "api": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3,  "cve": 4, "recommendation": 5, "new_obs": 6},
            "android": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3, "cve": 4, "recommendation": 5, "new_obs": 6},
            "ios": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3, "cve": 4, "recommendation": 5, "new_obs": 6},
            "sourcecode": {"sr_no": 0, "title": 1, "severity": 2, "affected": 3, "cve": 4, "recommendation": 5, "new_obs": 6}
        }
        
        print(f"DEBUG: template_type passed from GUI = '{template_type}'")

        # Define expected column count based on template type
        EXPECTED_COLUMNS = {
            "web": 7,
            "va": 6,
            "api": 7,
            "ca": 6,
            "ca_nessus": 7,
            "android": 7,
            "ios": 7,
            "sourcecode": 7,
        }

        expected_cols = EXPECTED_COLUMNS.get(template_type, 7)
        
        # Find the executive summary table (first table with headers)
        exec_table = None
        for table in doc.tables:
            if len(table.columns) == expected_cols and len(table.rows) > 0:
                header_text = " ".join([cell.text.strip().lower() for cell in table.rows[0].cells])
                # Check if it's the executive summary table (has "key issues" or "sr")
                if "key issues" in header_text or "sr" in header_text:
                    exec_table = table
                    break
        
        if not exec_table:
            print(f"⚠️  Executive summary table with {expected_cols} columns not found — skipping")
            return
        
        print(f"✅ Found {template_type.upper()} executive summary table with {expected_cols} columns")
        
        # Get mapping for this template type
        mapping = FIELD_MAPPING.get(template_type, FIELD_MAPPING["web"])
        print(f"DEBUG: mapping = {mapping}")
        
        # Clear all rows except header
        while len(exec_table.rows) > 1:
            exec_table._tbl.remove(exec_table.rows[1]._tr)
        
        # Severity colors
        severity_colors = {
            'high': 'EE0000',
            'critical': 'C00000',
            'medium': 'FFC000',
            'low': '00B050'
        }
        
        expected_cols = len(mapping)
        
        # Add observation rows
        for item in observations:
            cells = exec_table.add_row().cells
            
            # Build values array
            values = [""] * expected_cols
            
            values[mapping["sr_no"]] = item.get("sr_no", "")
            values[mapping["title"]] = item.get("title", "")
            values[mapping["severity"]] = item.get("severity", "")
            
            # Handle affected field based on template type
            if template_type == "ca_nessus":
                affected = item.get("host_ip", item.get("affected_ip", item.get("affected_url", "")))
            elif template_type == "va":
                affected = item.get("affected_ip", item.get("affected_url", ""))
            elif template_type == "api":
                affected = item.get("affected_endpoint", item.get("affected_url", ""))
            else:
                affected = item.get("affected_url", "")
            values[mapping["affected"]] = affected
            
            # Handle CVE field (if exists in this template)
            if "cve" in mapping:
                cve_val = item.get("cve", "")
                values[mapping["cve"]] = cve_val if cve_val else "Not Available"

            if "checkid" in mapping:
                checkid_val = item.get("checkid", "")
                values[mapping["checkid"]] = checkid_val if checkid_val else "Not Available"

            values[mapping["recommendation"]] = item.get("recommendation", "")
            values[mapping["new_obs"]] = "New Observation"
            
            # Apply formatting to each cell
            for idx, val in enumerate(values):
                cell = cells[idx]
                for para in cell.paragraphs:
                    para.clear()
                para = cell.paragraphs[0]
                text = str(val) if pd.notna(val) else ""
                run = para.add_run(text)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                if idx == mapping["severity"]:
                    severity_text = text.lower()
                    for sev, color in severity_colors.items():
                        if sev in severity_text:
                            run.font.bold = True
                            shading = parse_xml(
                                f'<w:shd xmlns:w="http://schemas.openxmlformats.org/'
                                f'wordprocessingml/2006/main" w:fill="{color}"/>'
                            )
                            cell._tc.get_or_add_tcPr().append(shading)
                            break
                
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == mapping["affected"] else WD_ALIGN_PARAGRAPH.JUSTIFY
        
        print(f"✅ Built {len(observations)} observations in executive summary")

    def _build_team_table(self, doc: Document, team_members: list[dict]):
        """
        Find the 6-column auditing team table and populate it.
        Columns: Sr.No | Name | Designation | Email | Qualifications | CERT-In Listed
        """
        team_table = next(
            (t for t in doc.tables
             if len(t.columns) == 6
             and "qualification" in t.rows[0].cells[4].text.strip().lower()),
            None
        )
        if not team_table:
                self._log("⚠️  Audit team table not found — skipping.")
                # Print all tables found for diagnosis
                for i, t in enumerate(doc.tables):
                    h = " | ".join(c.text.strip()[:15] for c in t.rows[0].cells)
                    self._log(f"    Table {i}: {len(t.columns)} cols → {h}")
                return
        self._log(f"✓ Audit team table found — adding {len(team_members)} members")

        # Clear all rows except header
        while len(team_table.rows) > 1:
            team_table._tbl.remove(team_table.rows[1]._tr)

        for i, member in enumerate(team_members, 1):
            cells  = team_table.add_row().cells
            values = [
                str(i),
                member.get("name", ""),
                member.get("designation", ""),
                member.get("email", ""),
                member.get("qualifications", ""),
                member.get("cert_in_listed", "Yes"),
            ]
            for idx, val in enumerate(values):
                cell = cells[idx]
                for p in cell.paragraphs:
                    p.clear()
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1
                run  = para.add_run(str(val))
                run.font.size      = Pt(10)
                run.font.name      = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 0)
                para.alignment     = WD_ALIGN_PARAGRAPH.CENTER


    def _build_tools_table(self, doc: Document, tools: list[dict]):
        """
        Find the 4-column tools table and populate it.
        Columns: Sr.No | Tool Name | Version | Open Source / Licensed
        """
        # Find 4-column table that is NOT the exec summary (which is 7 cols)
        tools_table = next(
            (t for t in doc.tables
             if len(t.columns) == 4
             and "name of tool" in t.rows[0].cells[1].text.strip().lower()),
            None
        )
        if not tools_table:
            self._log("⚠️  Tools table not found — skipping.")
            return

        # Clear all rows except header
        while len(tools_table.rows) > 1:
            tools_table._tbl.remove(tools_table.rows[1]._tr)

        for i, tool in enumerate(tools, 1):
            cells  = tools_table.add_row().cells
            values = [
                str(i),
                tool.get("tool_name", ""),
                tool.get("tool_version", ""),
                tool.get("tool_type", ""),
            ]
            for idx, val in enumerate(values):
                cell = cells[idx]
                for p in cell.paragraphs:
                    p.clear()
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1
                run  = para.add_run(str(val))
                run.font.size      = Pt(10)
                run.font.name      = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 0)
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2, 3]
                    else WD_ALIGN_PARAGRAPH.LEFT
                )

    def _build_scope_table(self, doc: Document, scope_items: list, app_type: str = "", environment: str = ""):
        """
        Find scope table by detecting 'Sr. No.' and 'Environment' headers.
        Works for any column count (3 or 4 columns).
        """
        scope_table = None
        for table in doc.tables:
            if len(table.columns) >= 2 and len(table.rows) > 0:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if any(term in headers[0] for term in ["sr", "s.no", "#", "sl no"]):
                    if any(term in headers[-1] for term in ["environment", "env"]):
                        scope_table = table
                        break

        if not scope_table:
            self._log("⚠️  Scope table not found — skipping.")
            return

        while len(scope_table.rows) > 1:
            scope_table._tbl.remove(scope_table.rows[1]._tr)

        num_cols = len(scope_table.columns)

        for i, item in enumerate(scope_items, 1):
            if item and item.strip():
                cells = scope_table.add_row().cells
                for idx, val in enumerate([str(i), str(item).strip()]):
                    cell = cells[idx]
                    for p in cell.paragraphs:
                        p.clear()
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1
                    run = para.add_run(val)
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if num_cols >= 3:
                    cell = cells[2]
                    for p in cell.paragraphs:
                        p.clear()
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1
                    run = para.add_run(app_type if num_cols == 4 else environment)
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if num_cols == 4:
                    cell = cells[3]
                    for p in cell.paragraphs:
                        p.clear()
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1
                    run = para.add_run(environment)
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._log(f"✓ Scope table populated with {len(scope_items)} items")

    @staticmethod
    def _remove_empty_rows(doc: Document):
        for table in doc.tables:
            empty = [i for i, row in enumerate(table.rows)
                     if all(c.text.strip() == "" for c in row.cells)]
            for i in reversed(empty):
                table._tbl.remove(table.rows[i]._tr)

    @staticmethod
    def _remove_jinja_artifact_paragraphs(doc: Document):
        """
        Remove ONLY paragraphs that consist entirely of a known Jinja
        control-tag remnant. We check the paragraph's raw XML for leftover
        markers — this is far safer than guessing based on "looks empty",
        which can accidentally delete structurally important paragraphs.
        """
        JINJA_MARKERS = ("{%", "%}", "{{", "}}")
        for para in list(doc.paragraphs):
            text = para.text.strip()
            # Only remove if paragraph text is EXACTLY a leftover Jinja
            # artifact (this should be empty after render, but if docxtpl
            # left any literal tag text behind due to template formatting
            # issues, clean those specific cases — never touch real content)
            if text == "" and not para.runs:
                # Genuinely empty AND has zero runs (no formatting objects,
                # no bookmarks, no field codes) — safe to remove
                if not para._element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                ) and not para._element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'
                ) and not para._element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart'
                ):
                    para._element.getparent().remove(para._element)
    
    @staticmethod
    def _update_chart_data(docx_path: str,  critical, high, medium, low):
        """
        Update the embedded vulnerability-distribution chart (Critical/High/Medium/Low
        bar chart) inside a saved .docx file — both the cached chart values
        and the embedded mini-Excel sheet behind "Edit Data".
        """
        values = {
            "Critical": int(critical) if str(critical).strip() else 0,
            "High":   int(high)   if str(high).strip()   else 0,
            "Medium": int(medium) if str(medium).strip() else 0,
            "Low":    int(low)    if str(low).strip()    else 0,
        }
 
        ns = {
            'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
 
        tmp_dir = docx_path + "_chart_tmp"
        if os.path.exists(tmp_dir):
            shutil.rmtree(tmp_dir)
        os.makedirs(tmp_dir)
 
        with zipfile.ZipFile(docx_path) as z:
            z.extractall(tmp_dir)
 
        chart_dir = os.path.join(tmp_dir, "word", "charts")
        if not os.path.isdir(chart_dir):
            shutil.rmtree(tmp_dir)
            return  # no chart in this template — nothing to do
 
        chart_files = [f for f in os.listdir(chart_dir)
                        if f.startswith("chart") and f.endswith(".xml")]
 
        for chart_file in chart_files:
            chart_path = os.path.join(chart_dir, chart_file)
            parser = etree.XMLParser(huge_tree=True)
            tree = etree.parse(chart_path, parser)
            root = tree.getroot()
            
            # Fix ptCount to match actual number of categories
            
 
            for ser in root.findall('.//c:ser', ns):
                for ptCount in ser.findall('.//c:ptCount', ns):
                    ptCount.set('val', str(len(values)))
                cat_pts = ser.findall('.//c:cat//c:strCache/c:pt', ns) \
                    or ser.findall('.//c:cat//c:numCache/c:pt', ns)
                categories = {}
                for pt in cat_pts:
                    idx = int(pt.get('idx'))
                    v = pt.find('c:v', ns)
                    if v is not None:
                        categories[idx] = v.text
 
                for pt in ser.findall('.//c:val//c:numCache/c:pt', ns):
                    idx = int(pt.get('idx'))
                    cat_name = categories.get(idx)
                    if cat_name in values:
                        v_elem = pt.find('c:v', ns)
                        v_elem.text = str(values[cat_name])

            # Fix axis scale to match new max value
            max_val = max(values.values()) if values else 10
            # Round up to a clean major unit
            import math
            if max_val <= 10:
                major_unit = 1
            elif max_val <= 50:
                major_unit = 5
            elif max_val <= 100:
                major_unit = 10
            elif max_val <= 500:
                major_unit = 50
            else:
                major_unit = int(math.ceil(max_val / 10 / 10) * 10)

            max_scale = major_unit * math.ceil((max_val + 1) / major_unit)

            # Update value axis (c:valAx) scaling
            for valAx in root.findall('.//c:valAx', ns):
                # Set max scale
                scaling = valAx.find('c:scaling', ns)
                if scaling is not None:
                    # Remove hardcoded max so Word auto-scales, OR set it explicitly
                    max_elem = scaling.find('c:max', ns)
                    if max_elem is not None:
                        max_elem.set('val', str(float(max_scale)))
                    min_elem = scaling.find('c:min', ns)
                    if min_elem is not None:
                        min_elem.set('val', '0.0')

                # Set major unit
                majorUnit = valAx.find('c:majorUnit', ns)
                if majorUnit is not None:
                    majorUnit.set('val', str(float(major_unit)))
 
            tree.write(chart_path, xml_declaration=True,
                       encoding="UTF-8", standalone=True)
 
            # Update embedded xlsx behind "Edit Data"
            rels_path = os.path.join(chart_dir, "_rels", chart_file + ".rels")
            if os.path.exists(rels_path):
                rels_tree = etree.parse(rels_path, etree.XMLParser(huge_tree=True))
                for rel in rels_tree.getroot():
                    target = rel.get('Target')
                    if target and target.lower().endswith('.xlsx'):
                        xlsx_path = os.path.normpath(os.path.join(chart_dir, target))
                        wb = openpyxl.load_workbook(xlsx_path)
                        ws = wb.active
                        for row in ws.iter_rows(min_row=2):
                            cat_name = row[0].value
                            if cat_name in values:
                                row[1].value = values[cat_name]
                        wb.save(xlsx_path)
 
        # Repackage docx in place
        repacked = docx_path + ".tmp"
        with zipfile.ZipFile(repacked, "w", zipfile.ZIP_DEFLATED) as zf:
            for root_dir, _, files in os.walk(tmp_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmp_dir)
                    zf.write(file_path, arcname)
 
        shutil.rmtree(tmp_dir)
        os.replace(repacked, docx_path)

    @staticmethod
    def _insert_page_breaks(doc: Document, observations: list[dict]):
        """
        Insert a real page break after each observation's marker,
        except after the last observation.
        """
        marker = "<!-- OBS_END -->"
        obs_index = 0
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if marker in para.text:
                is_last = (obs_index == len(observations) - 1)
                if is_last:
                    # Remove marker paragraph, no page break needed
                    para._element.getparent().remove(para._element)
                else:
                    # Replace marker text with a page break run
                    para.clear()
                    para.add_run().add_break(WD_BREAK.PAGE)
                    i += 1
                obs_index += 1
            else:
                i += 1

   

    def _insert_pocs(self, doc: Document, observations: list[dict], poc_finder: POCFinder):
        marker = "<!-- POC will be inserted here during post-processing -->"
        obs_index = 0
        i = 0

        MAX_HEIGHT_CM = 13
        MAX_WIDTH_CM = 21
        MAX_HEIGHT_INCHES = MAX_HEIGHT_CM / 2.54
        MAX_WIDTH_INCHES = MAX_WIDTH_CM / 2.54

        from docx.shared import Inches, Pt
        from docx.oxml import parse_xml
        from PIL import Image
        import os

        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if marker in para.text and obs_index < len(observations):
                vuln_title = observations[obs_index].get("title", "").strip()
                para._element.getparent().remove(para._element)

                if vuln_title:
                    # Get ordered items from Excel
                    ordered_items = poc_finder.get_excel_poc_items_by_vulnerability(vuln_title)
                    
                    # If no Excel POCs, try folder-based
                    if not ordered_items:
                        folder_images = poc_finder.get_pocs_by_vulnerability(vuln_title)
                        for img_path in folder_images:
                            ordered_items.append(('image', img_path))
                    
                    step_counter = 1
                    
                    for item in ordered_items:
                        item_type = item[0]
                        item_data = item[1]
                        
                        if item_type == 'step':
                            p = doc.paragraphs[i].insert_paragraph_before()
                            step_text = item_data
                            if not step_text.lower().startswith("step"):
                                run = p.add_run(f"Step {step_counter}: {step_text}")
                                step_counter += 1
                            else:
                                run = p.add_run(step_text)
                            run.font.name = "Calibri"
                            run.font.size = Pt(10)
                            i += 1
                        
                        elif item_type == 'image':
                            img_path = item_data
                            if os.path.exists(img_path):
                                try:
                                    with Image.open(img_path) as img:
                                        orig_width, orig_height = img.size

                                    orig_width_inches = orig_width / 96
                                    orig_height_inches = orig_height / 96

                                    needs_resize = (
                                        orig_width_inches > MAX_WIDTH_INCHES
                                        or orig_height_inches > MAX_HEIGHT_INCHES
                                    )

                                    p = doc.paragraphs[i].insert_paragraph_before()
                                    run = p.add_run()

                                    if needs_resize:
                                        width_ratio = MAX_WIDTH_INCHES / orig_width_inches
                                        height_ratio = MAX_HEIGHT_INCHES / orig_height_inches
                                        scale = min(width_ratio, height_ratio)
                                        new_width_inches = orig_width_inches * scale
                                        new_height_inches = orig_height_inches * scale
                                        run.add_picture(
                                            img_path,
                                            width=Inches(new_width_inches),
                                            height=Inches(new_height_inches),
                                        )
                                    else:
                                        run.add_picture(img_path)

                                    try:
                                        bdr_xml = (
                                            '<w:bdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                            'w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                                        )
                                        run._element.get_or_add_rPr().append(parse_xml(bdr_xml))
                                    except Exception as e:
                                        self._log(f"    ⚠️  Could not add border to {img_path}: {e}")

                                    i += 1

                                except Exception as e:
                                    self._log(f"    ⚠️  Could not add image {img_path}: {e}")
                                    p = doc.paragraphs[i].insert_paragraph_before()
                                    p.add_run(f"[POC image: {os.path.basename(img_path)}]")
                                    i += 1
                            else:
                                self._log(f"    ⚠️  Image not found: {img_path}")

                    if ordered_items:
                        self._log(f"    ✓ POCs added: {vuln_title[:40]}")
                    else:
                        self._log(f"    ⚠️  No POCs for: {vuln_title[:40]}")
                    
                    obs_index += 1
                else:
                    obs_index += 1
            else:
                i += 1
